#!/usr/bin/env python3
"""Generate Word thesis from JSON content data."""
import json, os, sys
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

FIGURES = r"d:\bylw\code\lunwen\QLULatex\QLUThesisLatexTemplate-master\Thesis\static\figures"
OUTPUT = r"d:\bylw\code\output_thesis.docx"

def R(para, text, fn, fs=Pt(12), bold=False):
    run = para.add_run(text)
    run.font.name = fn
    run._element.rPr.rFonts.set(qn('w:eastAsia'), fn)
    run.font.size = fs
    run.bold = bold
    return run

def P(doc, text='', fn='SimSun', fs=Pt(12), bold=False, align=None, indent=Cm(0.74), spacing=1.5):
    para = doc.add_paragraph()
    if align is not None:
        para.alignment = align
    fmt = para.paragraph_format
    fmt.line_spacing = spacing
    if indent:
        fmt.first_line_indent = indent
    if text:
        R(para, text, fn, fs, bold)
    return para

def H(doc, text, level=1):
    if level == 1:
        return P(doc, text, 'SimHei', Pt(16), True, WD_ALIGN_PARAGRAPH.CENTER, None)
    elif level == 2:
        return P(doc, text, 'SimHei', Pt(14), True, indent=None)
    else:
        return P(doc, text, 'SimHei', Pt(12), True, indent=None)

def B(doc, text):
    return P(doc, text, 'SimSun', Pt(12))

def FIG(doc, fname, caption, w=Inches(4.5)):
    path = os.path.join(FIGURES, fname)
    if os.path.exists(path):
        p = P(doc, '', indent=None, align=WD_ALIGN_PARAGRAPH.CENTER)
        p.add_run().add_picture(path, width=w)
        P(doc, caption, 'SimSun', Pt(10.5), indent=None, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    else:
        P(doc, '[MISSING: ' + fname + ']', 'SimSun', Pt(10.5), indent=None)

def TBL(doc, headers, rows, caption=''):
    if caption:
        P(doc, caption, 'SimSun', Pt(10.5), indent=None, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    for j, h in enumerate(headers):
        c = table.rows[0].cells[j]; c.text = ''
        R(c.paragraphs[0], h, 'SimHei', Pt(10.5), True)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = table.rows[i+1].cells[j]; c.text = ''
            R(c.paragraphs[0], str(val), 'SimSun', Pt(10.5))
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    P(doc, '', fs=Pt(6), indent=None)

def setup_page(section):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.2)

def build_cover(doc, m):
    for _ in range(3): P(doc, '', fs=Pt(16), indent=None)
    P(doc, m['cover_line1'], 'KaiTi', Pt(36), True, WD_ALIGN_PARAGRAPH.CENTER, None)
    P(doc, m['cover_line2'], 'KaiTi', Pt(36), True, WD_ALIGN_PARAGRAPH.CENTER, None)
    for _ in range(2): P(doc, '', fs=Pt(16), indent=None)
    P(doc, m['title_label'] + m['title'], 'SimHei', Pt(16), True, WD_ALIGN_PARAGRAPH.CENTER, None)
    for _ in range(2): P(doc, '', fs=Pt(14), indent=None)
    for key in ['school_label', 'class_label', 'author_label', 'id_label', 'supervisor_label']:
        P(doc, m[key], 'SimSun', Pt(14), indent=None, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(2): P(doc, '', fs=Pt(14), indent=None)
    P(doc, m['date'], 'SimSun', Pt(14), indent=None, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

def build_statement(doc, m):
    P(doc, m['stmt_title1'], 'SimHei', Pt(16), True, WD_ALIGN_PARAGRAPH.CENTER, None)
    P(doc, m['stmt_title2'], 'SimHei', Pt(16), True, WD_ALIGN_PARAGRAPH.CENTER, None)
    P(doc, '', fs=Pt(12), indent=None)
    B(doc, m['stmt_body'])
    P(doc, '', fs=Pt(12), indent=None)
    P(doc, m['stmt_sign'], indent=None)
    doc.add_page_break()

def build_abstract(doc, m):
    P(doc, m['cn_abs_title'], 'SimHei', Pt(16), True, WD_ALIGN_PARAGRAPH.CENTER, None)
    P(doc, '', fs=Pt(6), indent=None)
    for para in m['cn_abstract'].split('\n\n'):
        if para.strip(): B(doc, para.strip())
    P(doc, '', fs=Pt(6), indent=None)
    B(doc, m['cn_keywords_label'] + m['cn_keywords'])
    doc.add_page_break()
    P(doc, m['en_abs_title'], 'Times New Roman', Pt(16), True, WD_ALIGN_PARAGRAPH.CENTER, None)
    P(doc, '', fs=Pt(6), indent=None)
    for para in m['en_abstract'].split('\n\n'):
        if para.strip(): P(doc, para.strip(), 'Times New Roman', Pt(12))
    P(doc, '', fs=Pt(6), indent=None)
    P(doc, m['en_keywords_label'] + m['en_keywords'], 'Times New Roman', Pt(12))
    doc.add_page_break()

def build_chapter(doc, ch):
    H(doc, ch['number'] + '  ' + ch['title'], 1)
    for sec in ch.get('sections', []):
        H(doc, sec['title'], sec['level'])
        for text in sec.get('body', []):
            B(doc, text)
    for fig in ch.get('figures', []):
        FIG(doc, fig['file'], fig['caption'])
    for tbl in ch.get('tables', []):
        TBL(doc, tbl['headers'], tbl['rows'], tbl.get('caption', ''))

def build_refs(doc, data):
    H(doc, data['title'], 1)
    for ref in data['items']:
        P(doc, ref, 'SimSun', Pt(10.5))

def build_thanks(doc, data):
    H(doc, data['title'], 1)
    for t in data['body']:
        B(doc, t)

def add_page_numbers(doc, header_text):
    for i, section in enumerate(doc.sections):
        if i >= 4:
            hdr = section.header
            hdr.is_linked_to_previous = False
            hp = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
            hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            R(hp, header_text, 'SimSun', Pt(9))
        ftr = section.footer
        ftr.is_linked_to_previous = False
        fp = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = fp.add_run()
        r1._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
        r2 = fp.add_run()
        r2._r.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
        r3 = fp.add_run()
        r3._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))

def main():
    with open(sys.argv[1], 'r', encoding='utf-8') as f:
        data = json.load(f)

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    for section in doc.sections:
        setup_page(section)

    build_cover(doc, data['meta'])
    build_statement(doc, data['meta'])
    build_abstract(doc, data['meta'])

    P(doc, data['toc_title'], 'SimHei', Pt(16), True, WD_ALIGN_PARAGRAPH.CENTER, None)
    P(doc, data['toc_hint'], 'SimSun', Pt(10.5), indent=None, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    for ch in data['chapters']:
        print(f"Building {ch['number']}...")
        build_chapter(doc, ch)

    build_refs(doc, data['references'])
    build_thanks(doc, data['thanks'])

    add_page_numbers(doc, data['header_text'])

    doc.save(OUTPUT)
    print(f"Done! Saved to: {OUTPUT}")

if __name__ == '__main__':
    main()
