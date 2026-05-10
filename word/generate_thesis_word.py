#!/usr/bin/env python3
"""Convert LaTeX thesis to Word document matching QLU template format."""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE_DIR = r"d:\bylw\code\lunwen\QLULatex\QLUThesisLatexTemplate-master\Thesis"
FIGURES_DIR = os.path.join(BASE_DIR, "static", "figures")
OUTPUT = r"d:\bylw\code\智能仓储自主移动机器人结构设计与控制.docx"

THESIS_TITLE = "智能仓储自主移动机器人结构设计与控制"
SCHOOL = "机械工程学院"
CLASS = "机器人(SI)22-2"
AUTHOR = "王子煜"
STUDENT_ID = "202201230042"
SUPERVISOR = "刘鹏博"
DATE = "2026 年 03 月 10 日"
YEAR = "2026"

def add_run(para, text, font_name='宋体', font_size=None, bold=False):
    run = para.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if font_size:
        run.font.size = font_size
    run.bold = bold
    return run

def add_para(doc, text, font_name='宋体', font_size=Pt(12), bold=False,
             alignment=None, indent=Cm(0.74), spacing=1.5):
    para = doc.add_paragraph()
    if alignment is not None:
        para.alignment = alignment
    fmt = para.paragraph_format
    fmt.line_spacing = spacing
    if indent:
        fmt.first_line_indent = indent
    if text:
        add_run(para, text, font_name, font_size, bold)
    return para

def heading(doc, text, level=1):
    if level == 1:
        return add_para(doc, text, '黑体', Pt(16), True, WD_ALIGN_PARAGRAPH.CENTER, None, 1.5)
    elif level == 2:
        return add_para(doc, text, '黑体', Pt(14), True, indent=None, spacing=1.5)
    else:
        return add_para(doc, text, '黑体', Pt(12), True, indent=None, spacing=1.5)

def body(doc, text):
    return add_para(doc, text, '宋体', Pt(12), False, indent=Cm(0.74), spacing=1.5)

def add_fig(doc, fname, caption, width=Inches(4.5)):
    path = os.path.join(FIGURES_DIR, fname)
    if os.path.exists(path):
        p = add_para(doc, '', alignment=WD_ALIGN_PARAGRAPH.CENTER, indent=None)
        p.add_run().add_picture(path, width=width)
        add_para(doc, caption, '宋体', Pt(10.5), False, WD_ALIGN_PARAGRAPH.CENTER, None, 1.0)
    else:
        add_para(doc, '[图片缺失: ' + fname + ']', '宋体', Pt(10.5), indent=None)

def add_tbl(doc, headers, rows, caption=''):
    if caption:
        add_para(doc, caption, '宋体', Pt(10.5), False, WD_ALIGN_PARAGRAPH.CENTER, None, 1.0)
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    for j, h in enumerate(headers):
        c = table.rows[0].cells[j]; c.text = ''
        add_run(c.paragraphs[0], h, '黑体', Pt(10.5), True)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = table.rows[i+1].cells[j]; c.text = ''
            add_run(c.paragraphs[0], str(val), '宋体', Pt(10.5))
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, '', font_size=Pt(6), indent=None)
    return table

print("Script loaded. Content modules will be in separate files.")
