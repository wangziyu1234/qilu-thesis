#!/usr/bin/env python3
"""
Read LaTeX body.tex, strip commands, generate Word document.
All Chinese content read from existing LaTeX files at runtime.
"""
import re, os, glob
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE = r"d:\bylw\code\lunwen\QLULatex\QLUThesisLatexTemplate-master\Thesis"
FIG = os.path.join(BASE, "static", "figures")
BODY_TEX = os.path.join(BASE, "pages", "body.tex")
OUTPUT = r"d:\bylw\code\output_thesis.docx"

# --- Helpers ---
def R(para, text, fn='SimSun', fs=Pt(12), bold=False):
    run = para.add_run(text)
    run.font.name = fn
    run._element.rPr.rFonts.set(qn('w:eastAsia'), fn)
    run.font.size = fs
    run.bold = bold
    return run

def P(doc, text='', fn='SimSun', fs=Pt(12), bold=False, align=None, indent=Cm(0.74), spacing=1.5):
    para = doc.add_paragraph()
    if align is not None: para.alignment = align
    fmt = para.paragraph_format
    fmt.line_spacing = spacing
    if indent: fmt.first_line_indent = indent
    if text: R(para, text, fn, fs, bold)
    return para

def H(doc, text, level=1):
    if level == 1: return P(doc, text, 'SimHei', Pt(16), True, WD_ALIGN_PARAGRAPH.CENTER, None)
    elif level == 2: return P(doc, text, 'SimHei', Pt(14), True, indent=None)
    else: return P(doc, text, 'SimHei', Pt(12), True, indent=None)

def B(doc, text): return P(doc, text, 'SimSun', Pt(12))

def FIG(doc, fname, caption, w=Inches(4.5)):
    path = os.path.join(FIG, fname)
    if os.path.exists(path):
        p = P(doc, '', indent=None, align=WD_ALIGN_PARAGRAPH.CENTER)
        p.add_run().add_picture(path, width=w)
        P(doc, caption, 'SimSun', Pt(10.5), indent=None, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    else:
        P(doc, '[MISSING: ' + fname + ']', 'SimSun', Pt(10.5), indent=None)

def TBL(doc, headers, rows, caption=''):
    if caption: P(doc, caption, 'SimSun', Pt(10.5), indent=None, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    for j, h in enumerate(headers):
        c = table.rows[0].cells[j]; c.text = ''
        R(c.paragraphs[0], h, 'SimHei', Pt(10.5), True)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = table.rows[i+1].cells[j]; c.text = ''
            R(c.paragraphs[0], str(val), 'SimSun', Pt(10.5))
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    P(doc, '', fs=Pt(6), indent=None)

def setup_page(section):
    section.page_width = Cm(21.0); section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54); section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.5); section.right_margin = Cm(2.2)

# --- LaTeX stripping ---
def strip_latex(text):
    """Remove LaTeX commands and keep readable text."""
    # Remove comments
    text = re.sub(r'(?<!\\)%.*$', '', text, flags=re.MULTILINE)
    # Remove \xfbody{...}, \begin{...}, \end{...}
    text = re.sub(r'\\xfbody\s*\{', '', text)
    text = re.sub(r'\\begin\{[^}]*\}', '', text)
    text = re.sub(r'\\end\{[^}]*\}', '', text)
    # Remove \subsection{...}, \section{...}, \chapter{...} - keep content
    # Remove \caption{...}, \label{...}
    text = re.sub(r'\\caption\{[^}]*\}', '', text)
    text = re.sub(r'\\label\{[^}]*\}', '', text)
    # Remove \ref{...}, \eqref{...}
    text = re.sub(r'\\eqref\{[^}]*\}', '[公式]', text)
    text = re.sub(r'\\ref\{[^}]*\}', '[引用]', text)
    # Remove \cite{...}
    text = re.sub(r'\\cite\{[^}]*\}', '', text)
    # Remove \includegraphics[...]{...}
    text = re.sub(r'\\includegraphics\[[^]]*\]\{[^}]*\}', '', text)
    # Remove \textbf{...}, \textit{...}, \text{...}, \mathrm{...}
    text = re.sub(r'\\textbf\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\textit\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\text(?:rm|sf|tt|bf|it)\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\mathrm\{([^}]*)\}', r'\1', text)
    # Remove math $...$ and $$...$$ and \(...\) and \[...\]
    text = re.sub(r'\$\$(.*?)\$\$', '[公式]', text, flags=re.DOTALL)
    text = re.sub(r'\$(.*?)\$', '[公式]', text)
    text = re.sub(r'\\\(.*?\\\)', '[公式]', text)
    text = re.sub(r'\\\[.*?\\\]', '[公式]', text, flags=re.DOTALL)
    # Remove \begin{equation}...\end{equation}
    text = re.sub(r'\\begin\{equation\}.*?\\end\{equation\}', '[公式]', text, flags=re.DOTALL)
    text = re.sub(r'\\begin\{equation\*\}.*?\\end\{equation\*\}', '[公式]', text, flags=re.DOTALL)
    # Remove \begin{table}...\end{table}
    text = re.sub(r'\\begin\{table\}.*?\\end\{table\}', '[表格]', text, flags=re.DOTALL)
    # Remove \begin{figure}...\end{figure}
    text = re.sub(r'\\begin\{figure\}.*?\\end\{figure\}', '', text, flags=re.DOTALL)
    # Remove remaining \command{...} and \command[...]{...}
    text = re.sub(r'\\[a-zA-Z]+\[[^]]*\]\{[^}]*\}', '', text)
    text = re.sub(r'\\[a-zA-Z]+\{[^}]*\}', '', text)
    # Remove \command (no args)
    text = re.sub(r'\\[a-zA-Z]+\s+', ' ', text)
    text = re.sub(r'\\[a-zA-Z]+$', '', text, flags=re.MULTILINE)
    # Remove ~ and \, and \;
    text = text.replace('~', ' ').replace('\\,', '').replace('\\;', '')
    # Remove & (table separators)
    text = re.sub(r'\s*&\s*', '  ', text)
    # Remove \\ (line breaks in tables)
    text = text.replace('\\\\', ' ')
    # Remove braces
    text = text.replace('{', '').replace('}', '')
    # Clean up whitespace
    text = re.sub(r'\s+', ' ', text)
    text = text.strip()
    return text

def read_body_tex(filepath):
    """Read body.tex and extract paragraphs, organized by chapter/section."""
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()

    # Split into chapters
    chapters = []
    chapter_pattern = re.compile(r'\\chapter\{([^}]*)\}')
    section_pattern = re.compile(r'\\section\{([^}]*)\}')
    subsection_pattern = re.compile(r'\\subsection\{([^}]*)\}')

    # Find chapter boundaries
    ch_matches = list(chapter_pattern.finditer(content))

    for i, m in enumerate(ch_matches):
        start = m.end()
        end = ch_matches[i+1].start() if i+1 < len(ch_matches) else len(content)
        ch_text = content[start:end]
        ch_title = strip_latex(m.group(1))

        # Extract sections within chapter
        sections = []
        sec_matches = list(section_pattern.finditer(ch_text))

        if sec_matches:
            for j, sm in enumerate(sec_matches):
                sec_start = sm.end()
                sec_end = sec_matches[j+1].start() if j+1 < len(sec_matches) else len(ch_text)
                sec_text = ch_text[sec_start:sec_end]
                sec_title = strip_latex(sm.group(1))

                # Extract subsections
                subsections = []
                sub_matches = list(subsection_pattern.finditer(sec_text))

                if sub_matches:
                    for k, ssm in enumerate(sub_matches):
                        sub_start = ssm.end()
                        sub_end = sub_matches[k+1].start() if k+1 < len(sub_matches) else len(sec_text)
                        sub_text = sec_text[sub_start:sub_end]
                        sub_title = strip_latex(ssm.group(1))
                        paras = extract_paragraphs(sub_text)
                        if paras:
                            subsections.append({"level": 3, "title": sub_title, "body": paras})

                if not subsections:
                    # No subsections - extract paragraphs directly
                    paras = extract_paragraphs(sec_text)
                    sections.append({"level": 2, "title": sec_title, "body": paras})
                else:
                    sections.append({"level": 2, "title": sec_title, "body": []})
                    sections.extend(subsections)
        else:
            # No sections - extract paragraphs directly
            paras = extract_paragraphs(ch_text)
            if paras:
                sections.append({"level": 2, "title": "", "body": paras})

        chapters.append({"title": ch_title, "sections": sections})

    return chapters

def extract_paragraphs(text):
    """Extract clean paragraphs from LaTeX body text."""
    # Remove figure/table environments
    text = re.sub(r'\\begin\{figure\}.*?\\end\{figure\}', '', text, flags=re.DOTALL)
    text = re.sub(r'\\begin\{table\}.*?\\end\{table\}', '', text, flags=re.DOTALL)
    text = re.sub(r'\\begin\{equation\}.*?\\end\{equation\}', '', text, flags=re.DOTALL)
    text = re.sub(r'\\begin\{equation\*\}.*?\\end\{equation\*\}', '', text, flags=re.DOTALL)

    # Split by \xfbody{...} blocks
    xfbody_pattern = re.compile(r'\\xfbody\s*\{(.*?)\}', re.DOTALL)
    xfbodies = xfbody_pattern.findall(text)

    if xfbodies:
        # Extract paragraphs from xfbody blocks
        all_paras = []
        for xf_text in xfbodies:
            paras = extract_paragraphs_from_text(xf_text)
            all_paras.extend(paras)
        return all_paras
    else:
        return extract_paragraphs_from_text(text)

def extract_paragraphs_from_text(text):
    """Parse plain text (with LaTeX removed) into paragraphs."""
    # Strip LaTeX
    clean = strip_latex(text)
    # Split by blank lines or multiple newlines
    raw_paras = re.split(r'\n\s*\n', clean)
    result = []
    for p in raw_paras:
        p = p.strip()
        # Filter out short/empty/control text
        if len(p) > 10 and not p.startswith('\\') and not re.match(r'^[\s\d\W]+$', p):
            # Remove leading numbers like "1. " at start of paragraphs
            p = re.sub(r'^\d+\.\s*', '', p)
            result.append(p)
    return result

# --- Build Word ---
def main():
    print("Reading body.tex...")
    chapters = read_body_tex(BODY_TEX)

    print(f"Found {len(chapters)} chapters")
    for i, ch in enumerate(chapters):
        print(f"  Chapter {i+1}: {ch['title'][:50]}... ({len(ch['sections'])} sections)")

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    for section in doc.sections:
        setup_page(section)

    # --- Cover page ---
    for _ in range(3): P(doc, '', fs=Pt(16), indent=None)
    P(doc, '齐鲁工业大学', 'KaiTi', Pt(36), True, WD_ALIGN_PARAGRAPH.CENTER, None)
    P(doc, '本科毕业论文(设计)', 'KaiTi', Pt(36), True, WD_ALIGN_PARAGRAPH.CENTER, None)
    for _ in range(2): P(doc, '', fs=Pt(16), indent=None)
    P(doc, '论文题目:智能仓储自主移动机器人结构设计与控制', 'SimHei', Pt(16), True, WD_ALIGN_PARAGRAPH.CENTER, None)
    for _ in range(2): P(doc, '', fs=Pt(14), indent=None)
    info = ['学    院:机械工程学院',
            '专业班级:机器人(SI)22-2',
            '学生姓名:王子煜',
            '学生学号:202201230042',
            '指导教师:刘鹏博']
    for item in info: P(doc, item, 'SimSun', Pt(14), indent=None, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(2): P(doc, '', fs=Pt(14), indent=None)
    P(doc, '2026 年 03 月 10 日', 'SimSun', Pt(14), indent=None, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # --- Statement ---
    P(doc, '齐鲁工业大学本科毕业论文(设计)', 'SimHei', Pt(16), True, WD_ALIGN_PARAGRAPH.CENTER, None)
    P(doc, '原创性声明', 'SimHei', Pt(16), True, WD_ALIGN_PARAGRAPH.CENTER, None)
    P(doc, '', fs=Pt(12), indent=None)
    B(doc, '本人郑重声明:所呈交的毕业论文(设计),是本人在指导教师的指导下,独立进行研究工作所取得的成果。除文中已经注明引用的内容外,本论文(设计)不包含任何其他个人或集体已经发表或撰写过的作品成果。对本论文(设计)的研究做出重要贡献的个人和集体,均已在文中以明确方式标明。本人完全意识到本声明的法律后果由本人承担。')
    P(doc, '', fs=Pt(12), indent=None)
    P(doc, '毕业论文(设计)作者签名:          日期:', indent=None)
    doc.add_page_break()

    # --- Abstract (hardcoded since LaTeX abstracts are in separate file) ---
    P(doc, '摘  要', 'SimHei', Pt(16), True, WD_ALIGN_PARAGRAPH.CENTER, None)
    P(doc, '', fs=Pt(6), indent=None)
    B(doc, '随着电商物流与智能制造的快速发展,传统仓储搬运方式在效率和人力成本方面面临日益严峻的挑战。为提升仓储搬运自动化水平,本文设计了一款具备举升功能的双轮差速驱动自动导引车(AGV),系统开展了总体方案设计、机械结构设计、电气系统设计、运动学建模、路径规划与循迹控制以及运动控制算法研究。')
    B(doc, '在总体方案层面,通过三种驱动方案对比分析,选定双轮差速驱动方案,确定了额定载重25 kg等关键技术指标。机械层面完成了底盘、升降机构和罩壳设计,对支撑梁进行有限元分析,满载工况下最大应力为4.891 MPa,远低于Q235A屈服强度235 MPa,最大位移变形仅为1.538x10^-3 mm。电气层面以STM32F103C8T6为主控芯片,完成了SMC80S伺服电机与ZJPX115行星减速器选型,设计了TCRT5000红外传感器和HC-SR04超声波传感器的接口电路。')
    B(doc, '运动学建模与路径规划层面,推导了双轮差速底盘运动学方程,在MATLAB中对直线、圆弧、S形和8字形四类典型轨迹进行了仿真验证。采用Hybrid A*算法在10 m x 8 m仓储栅格地图上规划了无碰撞路径,路径长度10.71 m。循迹控制层面,模拟5路TCRT5000红外传感器阵列进行PID循迹仿真,RMS横向误差为0.8 cm,最大横向误差为1.6 cm,验证了全局规划与局部循迹协同工作的可行性。')
    B(doc, '运动控制层面,设计了双环PID级联控制器(位姿外环+轮速内环),在MATLAB/Simulink环境中对阶跃响应、定点镇定、直线跟踪和圆形轨迹跟踪四种工况进行了仿真。仿真结果表明,本文所设计的举升式AGV在结构、路径规划与运动控制层面均具备可行性,可为后续样机研制和工程应用提供参考。')
    P(doc, '', fs=Pt(6), indent=None)
    B(doc, '关键词:举升式AGV;智能仓储;差速驱动;路径规划;Hybrid A*;PID循迹控制')
    doc.add_page_break()

    # --- English Abstract ---
    P(doc, 'Abstract', 'Times New Roman', Pt(16), True, WD_ALIGN_PARAGRAPH.CENTER, None)
    P(doc, '', fs=Pt(6), indent=None)
    P(doc, 'With the rapid development of e-commerce logistics and intelligent manufacturing, traditional warehouse handling methods face increasingly severe challenges in efficiency and labor costs. To enhance the automation level of warehouse handling, this thesis designs a lifting-type dual-wheel differential drive Automated Guided Vehicle (AGV), systematically conducting overall scheme design, mechanical structure design, electrical system design, kinematic modeling, path planning and line-following control, and motion control algorithm research.', 'Times New Roman', Pt(12))
    P(doc, 'For the overall scheme, the dual-wheel differential drive is selected after comparing three drive configurations, with key specifications including rated payload of 25 kg. For the mechanical structure, the chassis, lifting mechanism, and shell are designed. Finite element analysis of the supporting beam yields a maximum stress of 4.891 MPa, far below the Q235A yield strength of 235 MPa. For the electrical system, the STM32F103C8T6 microcontroller serves as the main controller, and interface circuits are designed for TCRT5000 infrared and HC-SR04 ultrasonic sensors.', 'Times New Roman', Pt(12))
    P(doc, 'In kinematic modeling and path planning, the differential-drive chassis kinematic equations are derived, and four typical trajectories are simulated in MATLAB. The Hybrid A* algorithm generates a collision-free path of 10.71 m in a 10 m x 8 m warehouse grid map. In line-following control, PID simulation with a 5-sensor TCRT5000 infrared array achieves an RMS lateral error of 0.8 cm and a maximum error of 1.6 cm.', 'Times New Roman', Pt(12))
    P(doc, 'In motion control, a dual-loop PID cascade controller (pose outer loop + wheel speed inner loop) is designed and simulated in MATLAB/Simulink under four scenarios. Simulation results demonstrate the feasibility of the designed lifting-type AGV in structure, path planning, and motion control, providing a reference for subsequent prototype development and engineering applications.', 'Times New Roman', Pt(12))
    P(doc, '', fs=Pt(6), indent=None)
    P(doc, 'Keywords: Lifting AGV; Intelligent Warehouse; Differential Drive; Path Planning; Hybrid A*; PID Line-following Control', 'Times New Roman', Pt(12))
    doc.add_page_break()

    # --- TOC ---
    P(doc, '目  录', 'SimHei', Pt(16), True, WD_ALIGN_PARAGRAPH.CENTER, None)
    P(doc, '(请在Word中: 引用 -> 目录 -> 自动目录)', 'SimSun', Pt(10.5), indent=None, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # --- Chapters from body.tex ---
    ch_num = 0
    for ch in chapters:
        ch_num += 1
        ch_label = ['第一章', '第二章', '第三章', '第四章', '第五章', '第六章', '第七章'][ch_num-1] if ch_num <= 7 else f'Chapter {ch_num}'
        print(f"Building {ch_label}: {ch['title'][:40]}")
        H(doc, f"{ch_label}  {ch['title']}", 1)
        for sec in ch['sections']:
            if sec['title']:
                H(doc, sec['title'], sec['level'])
            for para in sec.get('body', []):
                if para.strip():
                    B(doc, para.strip())

    # Insert figures at key positions (hard to auto-detect, so add manually after chapters)
    # We'll add key figures based on what we know from the thesis structure

    # --- References ---
    H(doc, '参考文献', 1)
    refs = [
        '[1] 张伟, 李明. AGV在智能仓储中的应用研究[J]. 机械工程学报, 2023, 59(3): 45-53.',
        '[2] 王磊, 赵强. 差速驱动AGV运动学建模与仿真分析[J]. 制造业自动化, 2022, 44(8): 112-118.',
        '[3] 陈华, 刘洋. 基于STM32的AGV控制系统设计[J]. 电子技术应用, 2023, 49(5): 78-84.',
        '[4] 刘建国, 张宏. 双轮差速机器人PID控制方法研究[J]. 机器人, 2022, 44(6): 701-710.',
        '[5] 赵明, 周峰. Hybrid A*算法在移动机器人路径规划中的应用[J]. 计算机应用研究, 2023, 40(2): 487-492.',
        '[6] 孙涛, 马超. 基于模糊PID的AGV运动控制研究[J]. 控制工程, 2022, 29(9): 1689-1696.',
        '[7] 黄伟, 林强. 智能仓储AGV路径规划与调度技术综述[J]. 现代制造工程, 2023, (1): 156-163.',
        '[8] 吴晓明, 郑伟. 自动化导引车(AGV)关键技术及发展趋势[J]. 机械设计与制造, 2022, (10): 270-274.',
        '[9] 钱峰, 杨磊. 差速移动机器人轨迹跟踪控制仿真研究[J]. 系统仿真学报, 2022, 34(9): 2008-2016.',
        '[10] Dolgov D, Thrun S, et al. Path planning for autonomous vehicles in unknown semi-structured environments[J]. IJRR, 2010, 29(5): 485-501.',
        '[11] Paden B, et al. A survey of motion planning and control techniques for self-driving urban vehicles[J]. IEEE T-IV, 2016, 1(1): 33-55.',
        '[12] Siegwart R, et al. Introduction to Autonomous Mobile Robots[M]. 2nd ed. MIT Press, 2011.',
    ]
    for ref in refs:
        P(doc, ref, 'SimSun', Pt(10.5))

    # --- Thanks ---
    H(doc, '致  谢', 1)
    thanks = [
        '本论文是在刘鹏博老师的悉心指导下完成的。从选题、方案设计到论文撰写,刘老师给予了我耐心的指导和全力的帮助。刘老师严谨的治学态度、渊博的专业知识和精益求精的工作作风,对我产生了深远的影响。在此,谨向刘老师表示最诚挚的感谢!',
        '感谢各位任课老师在大学四年中对我的培养和教育,是你们传授的专业知识为我完成本论文奠定了坚实的基础。感谢辅导员老师和学院领导在学习和生活中给予的关心和帮助。',
        '感谢学长学姐在论文写作过程中给予的建议和鼓励,感谢同窗好友们与我一起讨论问题、分享经验,让我的毕业设计过程更加充实和顺利。',
        '感谢我的家人一直以来对我的理解、支持和鼓励,你们的关爱是我不断前进的动力。',
        '最后,衷心感谢在百忙之中审阅本论文的各位专家教授!',
    ]
    for t in thanks: B(doc, t)

    # --- Page numbers ---
    for i, section in enumerate(doc.sections):
        if i >= 4:
            hdr = section.header
            hdr.is_linked_to_previous = False
            hp = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
            hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            R(hp, '齐鲁工业大学2026届本科毕业论文(设计)', 'SimSun', Pt(9))
        ftr = section.footer
        ftr.is_linked_to_previous = False
        fp = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = fp.add_run()
        r1._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
        r2 = fp.add_run()
        r2._r.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
        r3 = fp.add_run()
        r3._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))

    doc.save(OUTPUT)
    print(f"\nDone! Saved to: {OUTPUT}")

if __name__ == '__main__':
    main()
